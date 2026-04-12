"""Convert translated markdown to a DOCX file."""

import io
import re
from docx import Document
from docx.shared import Pt, Inches


def markdown_to_docx(md_text: str) -> bytes:
    """Convert markdown text to DOCX bytes.

    Handles: headings (##, ###), bold (**), numbered lists, bullet lists,
    horizontal rules (---), and regular paragraphs.
    """
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for line in md_text.split("\n"):
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            continue

        # Horizontal rule
        if re.match(r"^-{3,}$", stripped):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            # Add a thin border line via a run of underscores
            run = p.add_run("_" * 70)
            run.font.size = Pt(6)
            run.font.color.rgb = None
            continue

        # Headings
        heading_match = re.match(r"^(#{1,4})\s+(.*)", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            h = doc.add_heading(level=min(level, 4))
            _add_rich_text(h, text)
            continue

        # Bullet list
        if stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:]
            p = doc.add_paragraph(style="List Bullet")
            _add_rich_text(p, text)
            continue

        # Numbered list
        num_match = re.match(r"^(\d+)[.)]\s+(.*)", stripped)
        if num_match:
            text = num_match.group(2)
            p = doc.add_paragraph(style="List Number")
            _add_rich_text(p, text)
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        _add_rich_text(p, stripped)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_rich_text(paragraph, text: str) -> None:
    """Parse inline **bold** markers and add runs to the paragraph."""
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)
